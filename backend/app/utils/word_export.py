import os
from docx import Document
from docx.shared import Pt
from datetime import datetime

EXPORT_FOLDER = "exports"

os.makedirs(EXPORT_FOLDER, exist_ok=True)


def export_word(title: str, content: str, output_path: str = None):

    # If no output path is supplied (Export feature),
    # save inside exports folder.
    if output_path is None:

        filename = f"{title.replace(' ', '_')}.docx"

        output_path = os.path.join(
            EXPORT_FOLDER,
            filename,
        )

    doc = Document()

    doc.add_heading("ResearchMind AI", level=1)

    doc.add_heading(title, level=2)

    doc.add_paragraph(
        datetime.now().strftime(
            "Generated on %d %B %Y, %I:%M %p"
        )
    )

    doc.add_paragraph()

    for line in content.split("\n"):

        line = line.strip()

        if not line:
            continue

        if line.startswith("#"):

            doc.add_heading(
                line.replace("#", "").strip(),
                level=2,
            )

        else:

            p = doc.add_paragraph()

            run = p.add_run(line)

            run.font.size = Pt(11)

    doc.save(output_path)

    return output_path